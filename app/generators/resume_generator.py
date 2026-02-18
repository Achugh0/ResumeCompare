import openai
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class ATSResumeGenerator:
    """
    Generates ATS-friendly resumes in Word and PDF formats.
    Applies improvement suggestions to create an enhanced resume.
    """
    
    def __init__(self, config):
        self.openai_api_key = config.get("OPENAI_API_KEY")
        openai.api_key = self.openai_api_key
    
    def generate_improved_resume(self, resume_text, suggestions_data, analysis_data, jd_text):
        """
        Generate an improved resume by applying suggestions.
        
        Args:
            resume_text: Original resume text
            suggestions_data: Suggestions from ImprovementEngine
            analysis_data: Analysis results
            jd_text: Job description text
            
        Returns:
            dict: Structured resume content
        """
        try:
            suggestions = suggestions_data.get("suggestions", [])
            
            # Build suggestions summary for AI
            suggestions_text = "\n".join([
                f"{i+1}. {s['area']}: {s['what_to_change']}\n   Before: {s['before']}\n   After: {s['after']}"
                for i, s in enumerate(suggestions)
            ])
            
            prompt = f"""
You are an expert resume writer specializing in ATS-optimized resumes.

**TASK**: Transform this resume by applying the improvement suggestions while maintaining an ATS-friendly format. Use a professional, accomplishment-driven tone.

**ORIGINAL RESUME**:
{resume_text}

**JOB DESCRIPTION** (for context):
{jd_text[:1500]}

**IMPROVEMENT SUGGESTIONS TO APPLY**:
{suggestions_text}

**INSTRUCTIONS**:
1. **Name and Contact**: Extract accurately. Ensure name is exactly as it appears.
2. **Professional Summary**: Write a powerful 2-3 sentence summary using the "After" versions of suggestions where applicable.
3. **Work Experience**:
   - List each job with Title, Company, Location, and Dates.
   - For each job, provide a list of 3-5 concise bullet points.
   - **MANDATORY**: Use the "After" bullet points provided in the suggestions.
   - Quantify achievements using numbers, percentages, or currency.
4. **Skills**: Categorize into "Technical Skills" and "Core Competencies".
5. **Education & Certifications**: List clearly with years.

**OUTPUT FORMAT** (Strict JSON):
{{
  "contact": {{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "+1234567890",
    "location": "City, State"
  }},
  "summary": "Impactful summary paragraph...",
  "experience": [
    {{
      "title": "Job Title",
      "company": "Company Name",
      "location": "City, State",
      "dates": "Start - End",
      "achievements": [
        "Concise achievement starting with an action verb (metric included)",
        "Another specific achievement point..."
      ]
    }}
  ],
  "skills": {{
    "Technical": ["Python", "AWS", "SQL"],
    "Competencies": ["Project Management", "Stakeholder Engagement"]
  }},
  "education": [
    {{
      "degree": "Degree Name",
      "institution": "University Name",
      "year": "YYYY"
    }}
  ],
  "certifications": ["Cert 1", "Cert 2"]
}}

**CRITICAL**:
- Output ONLY valid JSON.
- DO NOT wrap the content in one big paragraph.
- DO NOT use generic placeholders like [Your Name].
- Keep it clean: no tables, no columns, no colors.
"""
            
            resp = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=3000,
            )
            
            content = resp.choices[0].message.content
            if content.startswith("```json"):
                content = content[7:]
            if content.endswith("```"):
                content = content[:-3]
            
            resume_data = json.loads(content.strip())
            resume_data["_is_demo"] = False
            return resume_data
            
        except openai.RateLimitError as e:
            print(f"[RESUME GEN] OpenAI Rate Limit Error: {e}")
            return self._generate_template_resume(resume_text, suggestions_data)
        except openai.AuthenticationError as e:
            print(f"[RESUME GEN] OpenAI Authentication Error: {e}")
            return self._generate_template_resume(resume_text, suggestions_data)
        except openai.APIConnectionError as e:
            print(f"[RESUME GEN] OpenAI Connection Error: {e}")
            return self._generate_template_resume(resume_text, suggestions_data)
        except json.JSONDecodeError as e:
            print(f"[RESUME GEN] JSON Decode Error: {e}. Response was: {content if 'content' in locals() else 'N/A'}")
            return self._generate_template_resume(resume_text, suggestions_data)
        except Exception as e:
            print(f"[RESUME GEN] Unexpected Error: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
            return self._generate_template_resume(resume_text, suggestions_data)
    
    def _generate_template_resume(self, resume_text, suggestions_data):
        """
        Generate a comprehensive mock resume by deeply extracting real content.
        Ensures the "Improved Resume" reflects as much of the original as possible.
        """
        import re
        
        # 1. Basic Identity & Contact
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', resume_text)
        phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', resume_text)
        
        lines = [l.strip() for l in resume_text.splitlines() if l.strip()]
        
        # Heuristic for Name
        name = "Your Name"
        for l in lines[:10]:
            clean_l = l.lower()
            if not any(x in clean_l for x in ['@', 'phone', 'email', 'linkedin', 'address', 'mobile', 'location', 'summary', 'proven', 'experience']) \
               and 2 < len(l) < 40:
                name = l
                break

        # 2. Section Splitting
        sections = {
            "summary": "",
            "experience": [],
            "skills": {"Technical": [], "Soft Skills": []},
            "education": [],
            "certifications": []
        }
        
        current_section = None
        current_text_block = []
        
        # Keywords to detect section starts
        section_keywords = {
            "experience": ["experience", "employment", "work history", "professional history"],
            "education": ["education", "academic", "university", "college"],
            "skills": ["skills", "competencies", "strengths", "expertise", "expertise & skills"],
            "certifications": ["certifications", "licenses", "courses"]
        }
        
        # Detect sections and split text
        for line in lines:
            line_lower = line.lower()
            found_header = False
            for section, keywords in section_keywords.items():
                if any(kw in line_lower for kw in keywords) and len(line) < 30:
                    current_section = section
                    found_header = True
                    break
            
            if found_header:
                continue
            
            if current_section == "experience":
                # Very simple heuristic: a line with years might be a header or end of one
                # If we see a bullet, it's an achievement
                is_bullet = bool(re.match(r'^[•\-\*]\s+', line))
                if not is_bullet and (re.search(r'\d{4}', line) or len(line) < 60):
                     # New job entry? (simple heuristic)
                     if not sections["experience"] or sections["experience"][-1]["achievements"]:
                        sections["experience"].append({
                            "title": line,
                            "company": "Organization", # We'll try to find company in next line
                            "location": "Location",
                            "dates": re.search(r'\d{4}.*\d{4}|\d{4}.*Present', line).group(0) if re.search(r'\d{4}.*\d{4}|\d{4}.*Present', line) else "Dates",
                            "achievements": []
                        })
                     else:
                        # Probably company name if title was just set
                        sections["experience"][-1]["company"] = line
                elif is_bullet and sections["experience"]:
                    sections["experience"][-1]["achievements"].append(re.sub(r'^[•\-\*]\s+', '', line))
            
            elif current_section == "skills":
                # Split by commas or bullets
                parts = re.split(r'[,•\-\*|]', line)
                for p in parts:
                    clean_p = p.strip()
                    if clean_p and len(clean_p) > 2:
                        sections["skills"]["Technical"].append(clean_p)
            
            elif current_section == "education":
                if re.search(r'\d{4}', line) or "university" in line_lower or "college" in line_lower or "degree" in line_lower:
                    sections["education"].append({
                        "degree": line,
                        "institution": "University Name",
                        "year": re.search(r'\d{4}', line).group(0) if re.search(r'\d{4}', line) else "Year"
                    })
            
            elif current_section == "certifications":
                sections["certifications"].append(line)
        
        # 3. Fallbacks and Improvements
        if not sections["experience"]:
            # If no sections found, just take a chunk of text
            sections["experience"] = [{
                "title": "Professional Role",
                "company": "Your Company",
                "achievements": [l for l in lines[5:15] if len(l) > 30]
            }]
        
        # Apply AI suggestions to the first experience entry if available
        sugs = suggestions_data.get("suggestions", [])
        if sections["experience"] and sugs:
            # Add improved bullets to the first job for demonstration
            improved_bullets = [s["after"] for s in sugs[:3]]
            sections["experience"][0]["achievements"] = improved_bullets + sections["experience"][0]["achievements"][:2]

        return {
            "contact": {
                "name": name,
                "email": email_match.group(0) if email_match else "your.email@example.com",
                "phone": phone_match.group(0) if phone_match else "+1 (555) 123-4567",
                "location": "City, State"
            },
            "summary": "Result-oriented professional with experience in technical execution and complex problem-solving. This summary has been automatically tailored to align with the provided job description.",
            "experience": sections["experience"][:5], # Limit to 5 jobs
            "skills": {
                "Technical": sections["skills"]["Technical"][:15],
                "Core Competencies": ["Strategy", "Operations", "Management"]
            },
            "education": sections["education"][:2],
            "certifications": sections["certifications"][:3],
            "_is_demo": True,
            "_demo_reason": "API Quota Exceeded - Deep Mock Applied"
        }
    
    def create_docx(self, resume_data, output_path):
        """
        Create a Word document (.docx) from resume data.
        
        Args:
            resume_data: Structured resume content
            output_path: Path to save the .docx file
            
        Returns:
            str: Path to created file
        """
        doc = Document()
        
        # Set document margins (1 inch all around)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # CONTACT INFORMATION
        contact = resume_data.get("contact", {})
        
        # Name (Large, Bold)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(contact.get("name", "Your Name"))
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.name = 'Calibri'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Details (Centered)
        contact_parts = []
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("location"):
            contact_parts.append(contact["location"])
        if contact.get("linkedin"):
            contact_parts.append(contact["linkedin"])
        
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)
        contact_para.runs[0].font.name = 'Calibri'
        
        doc.add_paragraph()  # Spacing
        
        # PROFESSIONAL SUMMARY
        if resume_data.get("summary"):
            self._add_section_header(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(resume_data["summary"])
            summary_para.runs[0].font.size = Pt(11)
            summary_para.runs[0].font.name = 'Calibri'
            doc.add_paragraph()  # Spacing
        
        # WORK EXPERIENCE
        experience = resume_data.get("experience", [])
        if experience:
            self._add_section_header(doc, "WORK EXPERIENCE")
            
            for exp in experience:
                # Job Title (Bold)
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                title_run = title_para.add_run(exp.get("title", "Job Title"))
                title_run.font.size = Pt(11)
                title_run.font.bold = True
                title_run.font.name = 'Calibri'
                
                # Company, Location, Dates
                details_para = doc.add_paragraph()
                details_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                details_para.paragraph_format.space_after = Pt(2)
                details_text = f"{exp.get('company', 'Company')} | {exp.get('location', 'Location')} | {exp.get('dates', 'Dates')}"
                details_run = details_para.add_run(details_text)
                details_run.font.size = Pt(10)
                details_run.font.italic = True
                details_run.font.name = 'Calibri'
                
                # Achievements (Bullet Points)
                achievements = exp.get("achievements", [])
                for achievement in achievements:
                    bullet_para = doc.add_paragraph(achievement, style='List Bullet')
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
                    bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if bullet_para.runs:
                        bullet_para.runs[0].font.size = Pt(10)
                        bullet_para.runs[0].font.name = 'Calibri'
                
                temp_spacer = doc.add_paragraph()
                temp_spacer.paragraph_format.space_after = Pt(6)
        
        # SKILLS & EXPERTISE
        skills = resume_data.get("skills", {})
        if skills:
            self._add_section_header(doc, "SKILLS & EXPERTISE")
            
            for category, skill_list in skills.items():
                if skill_list:
                    skills_para = doc.add_paragraph()
                    category_run = skills_para.add_run(f"{category}: ")
                    category_run.font.bold = True
                    category_run.font.size = Pt(11)
                    category_run.font.name = 'Calibri'
                    
                    skills_run = skills_para.add_run(", ".join(skill_list))
                    skills_run.font.size = Pt(11)
                    skills_run.font.name = 'Calibri'
            
            doc.add_paragraph()  # Spacing
        
        # EDUCATION
        education = resume_data.get("education", [])
        if education:
            self._add_section_header(doc, "EDUCATION")
            
            for edu in education:
                edu_para = doc.add_paragraph()
                degree_run = edu_para.add_run(edu.get("degree", "Degree"))
                degree_run.font.bold = True
                degree_run.font.size = Pt(11)
                degree_run.font.name = 'Calibri'
                
                edu_para.add_run(f" | {edu.get('institution', 'Institution')} | {edu.get('year', 'Year')}")
                edu_para.runs[1].font.size = Pt(11)
                edu_para.runs[1].font.name = 'Calibri'
                
                if edu.get("details"):
                    details_para = doc.add_paragraph(edu["details"])
                    details_para.runs[0].font.size = Pt(10)
                    details_para.runs[0].font.italic = True
                    details_para.runs[0].font.name = 'Calibri'
            
            doc.add_paragraph()  # Spacing
        
        # CERTIFICATIONS
        certifications = resume_data.get("certifications", [])
        if certifications:
            self._add_section_header(doc, "CERTIFICATIONS")
            
            for cert in certifications:
                cert_para = doc.add_paragraph(cert, style='List Bullet')
                cert_para.runs[0].font.size = Pt(11)
                cert_para.runs[0].font.name = 'Calibri'
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def _add_section_header(self, doc, text):
        """Add a formatted section header with a line separator."""
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_para.paragraph_format.space_before = Pt(12)
        header_para.paragraph_format.space_after = Pt(2)
        
        header_run = header_para.add_run(text)
        header_run.font.size = Pt(11)
        header_run.font.bold = True
        header_run.font.name = 'Calibri'
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add a manual separator line (underscores) as a backup for border-less text
        line_para = doc.add_paragraph("_" * 90)
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_para.paragraph_format.space_after = Pt(6)
        line_para.runs[0].font.size = Pt(6)
        line_para.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    
    def create_pdf(self, resume_data, output_path):
        """
        Create a high-quality PDF from resume data using ReportLab.
        
        Args:
            resume_data: Structured resume content
            output_path: Path to save the .pdf file
            
        Returns:
            str: Path to created file
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib import colors

        # Document setup
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        elements = []

        # Custom styles
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['Heading1'],
            fontSize=22,
            alignment=TA_LEFT, # Changed to Left Align
            spaceAfter=8,
            fontName='Helvetica-Bold'
        )
        
        contact_style = ParagraphStyle(
            'ContactStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_LEFT, # Changed to Left Align
            spaceAfter=14,
            fontName='Helvetica'
        )
        
        header_style = ParagraphStyle(
            'HeaderStyle',
            parent=styles['Heading2'],
            fontSize=12,
            alignment=TA_LEFT,
            spaceBefore=16,
            spaceAfter=2, # Reduced to make room for hr
            fontName='Helvetica-Bold',
            textColor=colors.black
        )
        
        summary_style = ParagraphStyle(
            'SummaryStyle',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        job_title_style = ParagraphStyle(
            'JobTitleStyle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            spaceBefore=8,
            alignment=TA_LEFT
        )
        
        job_details_style = ParagraphStyle(
            'JobDetailsStyle',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica-Oblique',
            textColor=colors.grey,
            alignment=TA_LEFT,
            spaceAfter=4
        )

        def add_hr(elements):
            elements.append(Paragraph("<hr color='silver' width='100%' size='1'/>", styles['Normal']))

        # 1. Contact Info
        contact = resume_data.get("contact", {})
        elements.append(Paragraph(contact.get("name", "Your Name").upper(), name_style))
        
        contact_parts = []
        if contact.get("email"): contact_parts.append(contact["email"])
        if contact.get("phone"): contact_parts.append(contact["phone"])
        if contact.get("location"): contact_parts.append(contact["location"])
        
        elements.append(Paragraph(" | ".join(contact_parts), contact_style))
        
        # 2. Professional Summary
        summary = resume_data.get("summary")
        if summary:
            elements.append(Paragraph("PROFESSIONAL SUMMARY", header_style))
            add_hr(elements)
            elements.append(Paragraph(summary, summary_style))

        # 3. Work Experience
        experience = resume_data.get("experience", [])
        if experience:
            elements.append(Paragraph("WORK EXPERIENCE", header_style))
            add_hr(elements)
            
            for exp in experience:
                elements.append(Paragraph(exp.get("title", "Job Title"), job_title_style))
                elements.append(Paragraph(
                    f"{exp.get('company', 'Company')} | {exp.get('location', 'Location')} | {exp.get('dates', 'Dates')}",
                    job_details_style
                ))
                
                achievements = exp.get("achievements", [])
                if achievements:
                    bullet_items = [ListItem(Paragraph(a, summary_style)) for a in achievements]
                    elements.append(ListFlowable(bullet_items, bulletType='bullet', leftIndent=20, spaceBefore=4))
                
                elements.append(Spacer(1, 10))

        # 4. Skills
        skills = resume_data.get("skills", {})
        if skills:
            elements.append(Paragraph("SKILLS & EXPERTISE", header_style))
            elements.append(Paragraph("<hr color='black' width='100%' size='1'/>", styles['Normal']))
            
            for category, items in skills.items():
                if items:
                    elements.append(Paragraph(f"<b>{category}:</b> {', '.join(items)}", summary_style))
            
            elements.append(Spacer(1, 10))

        # 5. Education
        education = resume_data.get("education", [])
        if education:
            elements.append(Paragraph("EDUCATION", header_style))
            elements.append(Paragraph("<hr color='black' width='100%' size='1'/>", styles['Normal']))
            
            for edu in education:
                elements.append(Paragraph(
                    f"<b>{edu.get('degree', 'Degree')}</b> | {edu.get('institution', 'Institution')} | {edu.get('year', 'Year')}",
                    summary_style
                ))
                if edu.get("details"):
                    elements.append(Paragraph(edu["details"], job_details_style))
            
            elements.append(Spacer(1, 10))

        # 6. Certifications
        certifications = resume_data.get("certifications", [])
        if certifications:
            elements.append(Paragraph("CERTIFICATIONS", header_style))
            elements.append(Paragraph("<hr color='black' width='100%' size='1'/>", styles['Normal']))
            
            for cert in certifications:
                elements.append(Paragraph(f"• {cert}", summary_style))

        # Build PDF
        doc.build(elements)
        return output_path
