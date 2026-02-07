from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
import os

logger = logging.getLogger(__name__)

class TemplateGenerator:
    @staticmethod
    def create_resume_template(output_path):
        try:
            doc = Document()

            title = doc.add_heading("RESUME TEMPLATE", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            instructions = doc.add_paragraph()
            instructions.add_run("Instructions: ").bold = True
            instructions.add_run(
                "Fill in all sections below. This template is optimized for ATS compatibility."
            )

            doc.add_paragraph()
            doc.add_heading("PERSONAL INFORMATION", 1)
            doc.add_paragraph("Full Name: ________________________________________________")
            doc.add_paragraph("Email: ________________________________________________")
            doc.add_paragraph("Phone: ________________________________________________")
            doc.add_paragraph("LinkedIn: ________________________________________________")
            doc.add_paragraph("Location: ________________________________________________")

            doc.add_paragraph()
            doc.add_heading("PROFESSIONAL SUMMARY", 1)
            doc.add_paragraph(
                "[Write 3-4 sentences summarizing your background, key skills, and target role.]"
            )

            doc.add_paragraph()
            doc.add_heading("TECHNICAL SKILLS", 1)
            doc.add_paragraph(
                "• Programming Languages: [e.g., Java, Python, JavaScript, C++]"
            )
            doc.add_paragraph(
                "• Frameworks & Tools: [e.g., Spring Boot, React, Docker]"
            )
            doc.add_paragraph("• Databases: [e.g., MySQL, PostgreSQL, MongoDB]")
            doc.add_paragraph("• Cloud Platforms: [e.g., AWS, Azure, GCP]")
            doc.add_paragraph("• Other Skills: [e.g., Agile, Git, CI/CD]")

            doc.add_paragraph()
            doc.add_heading("CERTIFICATIONS", 1)
            doc.add_paragraph("• [Certification] - [Organization] - [Year]")
            doc.add_paragraph("• [Certification] - [Organization] - [Year]")

            doc.add_paragraph()
            doc.add_heading("PROFESSIONAL EXPERIENCE", 1)
            for i in range(1, 3):
                p = doc.add_paragraph()
                p.add_run(f"Job Title {i}").bold = True
                p.add_run(" | Company Name | Location\n")
                p.add_run("[Start Month Year] – [End Month Year or Present]").italic = True
                doc.add_paragraph(
                    "• [Achievement/responsibility with quantifiable results]"
                )
                doc.add_paragraph(
                    "• [Achievement/responsibility with quantifiable results]"
                )
                doc.add_paragraph(
                    "• [Key project or initiative and its impact]"
                )
                doc.add_paragraph(
                    "• [Technical skills utilized or leadership demonstrated]"
                )
                doc.add_paragraph()

            doc.add_heading("KEY PROJECTS", 1)
            p = doc.add_paragraph()
            p.add_run("Project Name").bold = True
            p.add_run(" | [Client/Company] | [Year]\n")
            doc.add_paragraph("• [Brief project description]")
            doc.add_paragraph("• [Technologies used]")
            doc.add_paragraph("• [Your role and contributions]")
            doc.add_paragraph("• [Measurable outcomes or impact]")

            doc.add_paragraph()
            doc.add_heading("EDUCATION", 1)
            doc.add_paragraph("[Degree] in [Field]")
            doc.add_paragraph("[University] | [Location] | [Year]")

            doc.add_paragraph()
            doc.add_heading("ADDITIONAL INFORMATION", 1)
            doc.add_paragraph("• Languages: [e.g., English, Hindi]")
            doc.add_paragraph("• Professional Memberships: [e.g., IEEE, ACM]")

            doc.save(output_path)
            logger.info(f"Resume template created: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error creating resume template: {e}")
            return False

    @staticmethod
    def create_job_description_template(output_path):
        try:
            doc = Document()

            title = doc.add_heading("JOB DESCRIPTION TEMPLATE", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            instructions = doc.add_paragraph()
            instructions.add_run("Instructions: ").bold = True
            instructions.add_run(
                "Fill in all sections below with specific requirements."
            )

            doc.add_paragraph()
            doc.add_heading("BASIC INFORMATION", 1)
            doc.add_paragraph("Job Title: ________________________________________________")
            doc.add_paragraph("Department: ________________________________________________")
            doc.add_paragraph("Location: ________________________________________________")
            doc.add_paragraph("Employment Type: [Full-time / Part-time / Contract]")
            doc.add_paragraph("Experience Level: [Entry / Mid / Senior / Lead]")

            doc.add_paragraph()
            doc.add_heading("COMPANY OVERVIEW", 1)
            doc.add_paragraph(
                "[Describe your company, mission, culture, and industry.]"
            )

            doc.add_paragraph()
            doc.add_heading("ROLE SUMMARY", 1)
            doc.add_paragraph(
                "[Provide a 2-3 sentence overview of the role and impact.]"
            )

            doc.add_paragraph()
            doc.add_heading("KEY RESPONSIBILITIES", 1)
            doc.add_paragraph("• [Primary responsibility]")
            doc.add_paragraph("• [Secondary responsibility]")
            doc.add_paragraph("• [Team responsibilities]")
            doc.add_paragraph("• [Strategic responsibilities]")

            doc.add_paragraph()
            doc.add_heading("REQUIRED QUALIFICATIONS", 1)
            doc.add_paragraph("Education:")
            doc.add_paragraph(
                "• [Degree level and field, e.g., B.Tech in Computer Science]"
            )
            doc.add_paragraph("Experience:")
            doc.add_paragraph(
                "• [Years of experience, e.g., 3+ years in Java development]"
            )
            doc.add_paragraph("Technical Skills (Must-Have):")
            doc.add_paragraph("• [Technologies, frameworks, tools]")

            doc.add_paragraph()
            doc.add_heading("PREFERRED QUALIFICATIONS", 1)
            doc.add_paragraph("• [Nice-to-have certifications, skills, domains]")

            doc.add_paragraph()
            doc.add_heading("SOFT SKILLS & COMPETENCIES", 1)
            doc.add_paragraph("• [Communication, teamwork, problem solving]")

            doc.add_paragraph()
            doc.add_heading("WORK ENVIRONMENT & CULTURE", 1)
            doc.add_paragraph(
                "[Remote/hybrid/onsite, team size, values, and culture.]"
            )

            doc.add_paragraph()
            doc.add_heading("IMPORTANT KEYWORDS FOR ATS", 1)
            doc.add_paragraph(
                "[List critical keywords: Java, Spring Boot, Microservices, etc.]"
            )

            doc.save(output_path)
            logger.info(f"Job description template created: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error creating JD template: {e}")
            return False

    @classmethod
    def generate_all_templates(cls, downloads_folder):
        os.makedirs(downloads_folder, exist_ok=True)
        resume_path = os.path.join(downloads_folder, "resume_template.docx")
        jd_path = os.path.join(downloads_folder, "job_description_template.docx")
        r_ok = cls.create_resume_template(resume_path)
        j_ok = cls.create_job_description_template(jd_path)
        return r_ok and j_ok
